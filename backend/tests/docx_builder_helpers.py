import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# --- Color Constants Matching PDF Master Template ---
COLOR_PRIMARY_BLUE = RGBColor(30, 86, 208)      # #1E56D0 (Royal Blue Accent)
COLOR_NAVY_HEADER = RGBColor(15, 23, 42)        # #0F172A (Table Header / Heading)
COLOR_TEXT_MAIN = RGBColor(51, 65, 85)          # #334155 (Body Slate Text)
COLOR_MUTED = RGBColor(100, 116, 139)           # #64748B (Muted Subtext)
HEX_NAVY_HEADER = "0F172A"
HEX_LIGHT_BG = "F8FAFC"
HEX_BORDER = "CBD5E1"
HEX_DASHED_BORDER = "94A3B8"
HEX_GREEN_BG = "DCFCE7"
HEX_GREEN_TEXT = "166534"
HEX_BLUE_BAR = "1E56D0"

def set_cell_background(cell, hex_color):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    """
    kwargs: top, bottom, left, right
    values: dict(val='single'/'dashed', sz='12', color='CBD5E1', space='0')
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), edge_data.get('color', 'CBD5E1'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_header_footer(section, category_name, logo_path):
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    hp.paragraph_format.space_after = Pt(4)
    
    # Table in header: Left = Logo + "CODE MORPHICX", Right = Category
    htab = header.add_table(rows=1, cols=2, width=Inches(6.77))
    htab.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_left = htab.cell(0, 0)
    c_right = htab.cell(0, 1)
    c_left.width = Inches(3.38)
    c_right.width = Inches(3.38)
    
    p_left = c_left.paragraphs[0]
    p_left.paragraph_format.space_after = Pt(0)
    p_left.paragraph_format.space_before = Pt(0)
    if os.path.exists(logo_path):
        run_logo = p_left.add_run()
        run_logo.add_picture(logo_path, width=Inches(0.22))
        run_space = p_left.add_run("  ")
    run_brand = p_left.add_run("CODE MORPHICX")
    run_brand.bold = True
    run_brand.font.size = Pt(8.5)
    run_brand.font.color.rgb = COLOR_NAVY_HEADER
    
    p_right = c_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_after = Pt(0)
    p_right.paragraph_format.space_before = Pt(0)
    run_cat = p_right.add_run(category_name.upper())
    run_cat.font.size = Pt(8.5)
    run_cat.font.color.rgb = COLOR_MUTED
    
    # Bottom border rule in header
    p_rule = header.add_paragraph()
    p_rule.paragraph_format.space_before = Pt(2)
    p_rule.paragraph_format.space_after = Pt(0)
    r_rule = p_rule.add_run("―" * 68)
    r_rule.font.size = Pt(6)
    r_rule.font.color.rgb = RGBColor(226, 232, 240)
    
    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = ""
    
    ftab = footer.add_table(rows=1, cols=2, width=Inches(6.77))
    ftab.alignment = WD_TABLE_ALIGNMENT.CENTER
    fc_left = ftab.cell(0, 0)
    fc_right = ftab.cell(0, 1)
    fc_left.width = Inches(3.38)
    fc_right.width = Inches(3.38)
    
    fp_left = fc_left.paragraphs[0]
    fp_left.paragraph_format.space_after = Pt(0)
    fp_left.paragraph_format.space_before = Pt(0)
    r_fleft = fp_left.add_run("Code Morphicx")
    r_fleft.font.size = Pt(8.5)
    r_fleft.font.color.rgb = COLOR_MUTED
    
    fp_right = fc_right.paragraphs[0]
    fp_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp_right.paragraph_format.space_after = Pt(0)
    fp_right.paragraph_format.space_before = Pt(0)
    r_fright = fp_right.add_run("Project Documentation")
    r_fright.font.size = Pt(8.5)
    r_fright.font.color.rgb = COLOR_MUTED

def add_section_title(doc, number_str, title_str=None):
    if title_str is None:
        if '.' in number_str:
            parts = number_str.split('.', 1)
            number_str = parts[0].strip()
            title_str = parts[1].strip()
        else:
            title_str = number_str
            number_str = ""
            
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(10)
    
    # Blue vertical bar
    r_bar = p.add_run("|  ")
    r_bar.bold = True
    r_bar.font.size = Pt(18)
    r_bar.font.color.rgb = COLOR_PRIMARY_BLUE
    
    if number_str:
        r_num = p.add_run(f"{number_str}. ")
        r_num.bold = True
        r_num.font.size = Pt(18)
        r_num.font.color.rgb = COLOR_NAVY_HEADER
    
    r_title = p.add_run(title_str)
    r_title.bold = True
    r_title.font.size = Pt(18)
    r_title.font.color.rgb = COLOR_NAVY_HEADER

def add_subsection_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = COLOR_NAVY_HEADER

def add_body_paragraph(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = COLOR_TEXT_MAIN
    return p

def add_bullet_point(doc, text, space_after=3):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = COLOR_TEXT_MAIN
    return p

def add_status_badge(doc, text="PROJECT DOCUMENTATION"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    
    # Create single-cell table for pill badge
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    cell.width = Inches(2.2)
    set_cell_background(cell, HEX_GREEN_BG)
    set_cell_margins(cell, top=60, bottom=60, left=120, right=120)
    set_cell_border(cell, 
                    top=dict(val='single', sz=4, color='86EFAC'),
                    bottom=dict(val='single', sz=4, color='86EFAC'),
                    left=dict(val='single', sz=4, color='86EFAC'),
                    right=dict(val='single', sz=4, color='86EFAC'))
    
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(0)
    cp.paragraph_format.space_before = Pt(0)
    r = cp.add_run(text)
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(22, 101, 52) # Dark green

def add_card_box(doc, title, text, space_after=8):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.77)
    set_cell_background(cell, HEX_LIGHT_BG)
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    set_cell_border(cell, 
                    top=dict(val='single', sz=4, color=HEX_BORDER),
                    bottom=dict(val='single', sz=4, color=HEX_BORDER),
                    left=dict(val='single', sz=4, color=HEX_BORDER),
                    right=dict(val='single', sz=4, color=HEX_BORDER))
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r_t = p.add_run(f"{title}\n")
    r_t.bold = True
    r_t.font.size = Pt(10.5)
    r_t.font.color.rgb = COLOR_PRIMARY_BLUE
    
    r_b = p.add_run(text)
    r_b.font.size = Pt(9.5)
    r_b.font.color.rgb = COLOR_TEXT_MAIN
    
    # Extra paragraph after table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(space_after)

def add_dashed_placeholder_box(doc, text_content, height_inches=1.8, space_after=8):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(6.77)
    set_cell_background(cell, HEX_LIGHT_BG)
    set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
    set_cell_border(cell, 
                    top=dict(val='dashed', sz=6, color=HEX_DASHED_BORDER),
                    bottom=dict(val='dashed', sz=6, color=HEX_DASHED_BORDER),
                    left=dict(val='dashed', sz=6, color=HEX_DASHED_BORDER),
                    right=dict(val='dashed', sz=6, color=HEX_DASHED_BORDER))
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    
    lines = text_content.split('\n')
    for i, line in enumerate(lines):
        r = p.add_run(line)
        r.font.size = Pt(9.5)
        if "INSERT" in line or "DIAGRAM" in line or "WORKFLOW" in line or "SCREENSHOT" in line:
            r.bold = True
            r.font.color.rgb = COLOR_NAVY_HEADER
        else:
            r.font.color.rgb = COLOR_TEXT_MAIN
        if i < len(lines) - 1:
            p.add_run('\n')
            
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(space_after)

def add_styled_table(doc, headers, data_rows, col_widths=None):
    num_cols = len(headers)
    table = doc.add_table(rows=len(data_rows) + 1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header Row
    hdr_row = table.rows[0]
    for col_idx, header_text in enumerate(headers):
        cell = hdr_row.cells[col_idx]
        set_cell_background(cell, HEX_NAVY_HEADER)
        set_cell_margins(cell, top=90, bottom=90, left=120, right=120)
        set_cell_border(cell, 
                        top=dict(val='single', sz=4, color=HEX_NAVY_HEADER),
                        bottom=dict(val='single', sz=4, color=HEX_NAVY_HEADER),
                        left=dict(val='single', sz=4, color=HEX_NAVY_HEADER),
                        right=dict(val='single', sz=4, color=HEX_NAVY_HEADER))
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header_text)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    # Data Rows
    for row_idx, row_data in enumerate(data_rows):
        row = table.rows[row_idx + 1]
        bg_color = "FFFFFF" if row_idx % 2 == 0 else "F8FAFC"
        for col_idx, cell_value in enumerate(row_data):
            cell = row.cells[col_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            set_cell_border(cell, 
                            top=dict(val='single', sz=4, color="E2E8F0"),
                            bottom=dict(val='single', sz=4, color="E2E8F0"),
                            left=dict(val='single', sz=4, color="E2E8F0"),
                            right=dict(val='single', sz=4, color="E2E8F0"))
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(cell_value))
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0 and ("FR-" in str(cell_value) or "TC-" in str(cell_value) or str(cell_value).isdigit()):
                r.bold = True
            if str(cell_value) == "Pass":
                r.bold = True
                r.font.color.rgb = RGBColor(22, 101, 52)
            elif str(cell_value) == "High":
                r.bold = True
                r.font.color.rgb = RGBColor(185, 28, 28)
            elif str(cell_value) == "Medium":
                r.bold = True
                r.font.color.rgb = RGBColor(180, 83, 9)
                
    # Apply widths if provided
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)
                
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(2)
    p_sp.paragraph_format.space_after = Pt(8)

print("Helper functions ready!")
