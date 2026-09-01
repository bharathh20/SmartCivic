import os
import docx

doc_path = r"c:\Users\Bharath S\OneDrive\Desktop\smart civic\SmartCivic_Project_Documentation.docx"
doc = docx.Document(doc_path)

print(f"File Size: {os.path.getsize(doc_path)} bytes")
print(f"Total Sections: {len(doc.sections)}")
print(f"Total Paragraphs: {len(doc.paragraphs)}")
print(f"Total Tables: {len(doc.tables)}")

for idx, sec in enumerate(doc.sections, 1):
    header_text = ""
    if sec.header:
        for p in sec.header.paragraphs:
            if p.text.strip():
                header_text += p.text.strip() + " "
    safe_header = header_text.strip().encode('ascii', 'replace').decode('ascii')
    print(f"Section {idx}: Header='{safe_header}' | TopMargin={sec.top_margin.inches:.2f}in, LeftMargin={sec.left_margin.inches:.2f}in")
