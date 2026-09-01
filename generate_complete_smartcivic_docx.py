import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, parse_xml
sys.path.append(os.path.join(os.path.dirname(__file__), 'backend', 'tests'))

from docx_builder_helpers import (
    COLOR_PRIMARY_BLUE, COLOR_NAVY_HEADER, COLOR_TEXT_MAIN, COLOR_MUTED,
    HEX_NAVY_HEADER, HEX_LIGHT_BG, HEX_BORDER, HEX_DASHED_BORDER,
    HEX_GREEN_BG, HEX_GREEN_TEXT, HEX_BLUE_BAR,
    set_cell_background, set_cell_margins, set_cell_border,
    add_header_footer, add_section_title, add_subsection_heading,
    add_body_paragraph, add_bullet_point, add_status_badge,
    add_card_box, add_dashed_placeholder_box, add_styled_table
)

def build_smartcivic_documentation():
    print("Initializing SmartCivic 26-Page Master Documentation Generator...")
    doc = Document()
    
    # Configure Normal Style Font
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Segoe UI'
    font.size = Pt(10)
    font.color.rgb = COLOR_TEXT_MAIN

    logo_path = os.path.abspath(r"c:\Users\Bharath S\OneDrive\Desktop\smart civic\codemorphicx_logo.jpg")

    # =========================================================================
    # PAGE 1: COVER PAGE
    # =========================================================================
    sec1 = doc.sections[0]
    sec1.top_margin = Inches(0.6)
    sec1.bottom_margin = Inches(0.6)
    sec1.left_margin = Inches(0.75)
    sec1.right_margin = Inches(0.75)
    sec1.different_first_page_header_footer = True
    
    # Top Accent Blue Bar
    p_top_bar = doc.add_paragraph()
    p_top_bar.paragraph_format.space_before = Pt(0)
    p_top_bar.paragraph_format.space_after = Pt(24)
    r_bar = p_top_bar.add_run("━" * 58)
    r_bar.font.size = Pt(14)
    r_bar.font.color.rgb = COLOR_PRIMARY_BLUE
    
    # Code Morphicx Logo
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(10)
        p_logo.paragraph_format.space_after = Pt(12)
        p_logo.add_run().add_picture(logo_path, width=Inches(1.15))
        
    # Brand Name & Tagline
    p_brand = doc.add_paragraph()
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_brand.paragraph_format.space_before = Pt(0)
    p_brand.paragraph_format.space_after = Pt(2)
    r_b = p_brand.add_run("CODE MORPHICX")
    r_b.bold = True
    r_b.font.size = Pt(20)
    r_b.font.color.rgb = COLOR_NAVY_HEADER
    
    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tag.paragraph_format.space_before = Pt(0)
    p_tag.paragraph_format.space_after = Pt(40)
    r_t = p_tag.add_run("Transforming Ideas into Innovation")
    r_t.font.size = Pt(10.5)
    r_t.font.color.rgb = COLOR_MUTED
    
    # Document Title
    p_title1 = doc.add_paragraph()
    p_title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title1.paragraph_format.space_before = Pt(10)
    p_title1.paragraph_format.space_after = Pt(0)
    r_t1 = p_title1.add_run("PROJECT")
    r_t1.bold = True
    r_t1.font.size = Pt(24)
    r_t1.font.color.rgb = COLOR_NAVY_HEADER
    
    p_title2 = doc.add_paragraph()
    p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title2.paragraph_format.space_before = Pt(0)
    p_title2.paragraph_format.space_after = Pt(12)
    r_t2 = p_title2.add_run("DOCUMENTATION")
    r_t2.bold = True
    r_t2.font.size = Pt(24)
    r_t2.font.color.rgb = COLOR_NAVY_HEADER
    
    # Short Blue Accent Line
    p_sbar = doc.add_paragraph()
    p_sbar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sbar.paragraph_format.space_before = Pt(0)
    p_sbar.paragraph_format.space_after = Pt(14)
    r_sb = p_sbar.add_run("━━━━━━")
    r_sb.bold = True
    r_sb.font.size = Pt(12)
    r_sb.font.color.rgb = COLOR_PRIMARY_BLUE
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(45)
    r_sub = p_sub.add_run("4th Batch Internship Program")
    r_sub.bold = True
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = COLOR_PRIMARY_BLUE
    
    # Metadata Table
    meta_rows = [
        ["Project Title", "SmartCivic — Transparent Citizen Grievance & Municipal Resolution Platform"],
        ["Project Type", "Full Stack Web Application [FSD / MERN Architecture]"],
        ["Domain", "FSD / Web Development / Civic Technology"],
        ["Intern Name", "Bharath S"],
        ["Intern ID / USN", "[ID / USN]"],
        ["Internship Batch", "4th Batch"],
        ["Organization", "Code Morphicx"]
    ]
    
    meta_table = doc.add_table(rows=len(meta_rows), cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, (k, v) in enumerate(meta_rows):
        row = meta_table.rows[r_idx]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width = Inches(2.2)
        c2.width = Inches(4.57)
        set_cell_background(c1, "F8FAFC")
        set_cell_background(c2, "FFFFFF")
        set_cell_margins(c1, top=80, bottom=80, left=120, right=120)
        set_cell_margins(c2, top=80, bottom=80, left=120, right=120)
        set_cell_border(c1, top=dict(val='single', sz=4, color='E2E8F0'), bottom=dict(val='single', sz=4, color='E2E8F0'), left=dict(val='single', sz=4, color='E2E8F0'), right=dict(val='single', sz=4, color='E2E8F0'))
        set_cell_border(c2, top=dict(val='single', sz=4, color='E2E8F0'), bottom=dict(val='single', sz=4, color='E2E8F0'), left=dict(val='single', sz=4, color='E2E8F0'), right=dict(val='single', sz=4, color='E2E8F0'))
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(0); p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(k)
        r1.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = COLOR_NAVY_HEADER
        
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(v)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = COLOR_TEXT_MAIN
        
    p_dline = doc.add_paragraph()
    p_dline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dline.paragraph_format.space_before = Pt(50)
    p_dline.paragraph_format.space_after = Pt(0)
    r_dl = p_dline.add_run("Project Documentation Submission Deadline: 28 August 2026")
    r_dl.font.size = Pt(9)
    r_dl.font.color.rgb = COLOR_MUTED

    # Helper function for adding standard content pages
    def start_new_page(category_header_text):
        new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
        new_sec.top_margin = Inches(0.65)
        new_sec.bottom_margin = Inches(0.65)
        new_sec.left_margin = Inches(0.75)
        new_sec.right_margin = Inches(0.75)
        add_header_footer(new_sec, category_header_text, logo_path)
        return new_sec

    # =========================================================================
    # PAGE 2: 01. DOCUMENT CONTROL
    # =========================================================================
    start_new_page("PROJECT DOCUMENTATION")
    add_section_title(doc, "01", "Document Control")
    
    doc_ctrl_data = [
        ["Project Title", "SmartCivic — Transparent Citizen Grievance & Municipal Resolution Platform"],
        ["Document Title", "Project Documentation"],
        ["Document Version", "1.0"],
        ["Prepared By", "Bharath S"],
        ["Reviewed By", "Code Morphicx Technical Team"],
        ["Organization", "Code Morphicx"],
        ["Internship Batch", "4th Batch"],
        ["Submission Date", "28/08/2026"]
    ]
    add_styled_table(doc, ["Document Information", "Details"], doc_ctrl_data, col_widths=[2.5, 4.27])
    
    add_subsection_heading(doc, "Document Purpose")
    add_body_paragraph(doc, 
        "This document provides the complete functional, structural, and technical documentation of the SmartCivic platform assigned and developed as part of the Code Morphicx Internship Program. It details the end-to-end system architecture, comprehensive business requirements, user and administrative functional specifications, database models, RESTful API documentation, security implementations, automated testing suites, and cloud deployment procedures for academic review and operational evaluation.", space_after=14)
    
    add_subsection_heading(doc, "Document Status")
    add_status_badge(doc, "PROJECT DOCUMENTATION")

    # =========================================================================
    # PAGE 3: 02. PROJECT OVERVIEW
    # =========================================================================
    start_new_page("PROJECT OVERVIEW")
    add_section_title(doc, "02", "Project Overview")
    
    add_subsection_heading(doc, "2.1 Project Introduction")
    add_body_paragraph(doc, 
        "SmartCivic is a modern, transparent civic grievance reporting and municipal resolution platform engineered specifically for urban governance in Bengaluru Municipal Zone C. The platform empowers citizens to report localized public infrastructure failures—such as hazardous potholes, broken streetlights, overflowing waste bins, and water pipeline leakages—with photographic evidence, precise geolocation coordinates, and automated SLA tracking. The application connects citizens, municipal field dispatchers, and zonal administrative commissioners through a unified, high-performance web interface.")
    
    add_subsection_heading(doc, "2.2 Business Context")
    add_body_paragraph(doc, 
        "Rapid urbanization in metropolitan zones has placed tremendous strain on civic maintenance infrastructure. Traditional grievance channels suffer from opaque ticketing, slow departmental routing, and lack of visual proof, leading to citizen distrust and municipal dispatch bottlenecks. SmartCivic establishes an accountable digital ecosystem connecting citizens directly to urban local bodies (PWD, BESCOM, BBMP, BWSSB) for verifiable grievance resolution.")
    
    add_subsection_heading(doc, "2.3 Problem Statement")
    add_dashed_placeholder_box(doc, 
        "Urban citizens currently lack a centralized, transparent platform to report civic infrastructure defects with photographic evidence and track their resolution progress in real time. Municipal authorities struggle with manual grievance triage, duplicated complaint reports, and fragmented departmental dispatch workflows, leading to delayed SLA compliance, lack of accountability, and citizen disengagement.", height_inches=1.2)
    
    add_subsection_heading(doc, "2.4 Project Objectives")
    add_bullet_point(doc, "Deliver an intuitive, mobile-first web portal providing seamless grievance reporting with photo evidence upload and geolocation coordinates.")
    add_bullet_point(doc, "Implement strict citizen complaint ownership and isolation, ensuring citizens view and manage strictly their own grievances while admins access city-wide records.")
    add_bullet_point(doc, "Engineer an automated SLA triage engine that categorizes civic complaints, assigns priority ratings, and establishes verifiable resolution deadlines.")
    add_bullet_point(doc, "Equip municipal dispatch officers with dedicated department portals to inspect evidence, dispatch field crews, and append live timeline updates.")
    
    add_subsection_heading(doc, "2.5 Target Users")
    add_body_paragraph(doc, 
        "1. Urban Citizens: Residents seeking a reliable, transparent channel to report neighborhood issues and track resolution timelines.\n"
        "2. Department Officers & Field Engineers: Municipal personnel (PWD, BESCOM, BBMP Sanitation, BWSSB) responsible for inspecting, assigning crews, and resolving tickets.\n"
        "3. Municipal Administrators: Zonal commissioners requiring executive visibility, SLA compliance statistics, and city-wide grievance analytics.")
    
    add_subsection_heading(doc, "2.6 Project Scope")
    add_body_paragraph(doc, 
        "The scope includes JWT-based user authentication, a 3-step complaint wizard with evidence photo upload, isolated citizen dashboards, live milestone timelines, department dispatch triage, and single-service cloud deployment on Render with MongoDB Atlas.")

    # =========================================================================
    # PAGE 4: 03. BUSINESS REQUIREMENTS
    # =========================================================================
    start_new_page("BUSINESS REQUIREMENTS")
    add_section_title(doc, "03", "Business Requirements")
    
    add_subsection_heading(doc, "3.1 Business Objective")
    add_body_paragraph(doc, 
        "The primary business objective is to streamline urban municipal governance by reducing grievance turnaround times by 40%, elevating SLA resolution compliance above 85%, eliminating manual administrative overhead, and fostering citizen trust through end-to-end operational transparency.")
    
    add_subsection_heading(doc, "3.2 Business Requirements")
    add_bullet_point(doc, "Transparent Public Grievance Tracking: Enable public and authenticated tracking of civic issues with verifiable progress bars and audit notes.")
    add_bullet_point(doc, "Photographic Evidence Verification: Mandate photo attachment verification during complaint logging to prevent fraudulent or ambiguous submissions.")
    add_bullet_point(doc, "Automated SLA Triage & Department Routing: Automatically classify complaints into designated departments (PWD, BESCOM, BBMP, BWSSB) with target resolution deadlines.")
    add_bullet_point(doc, "Strict Role-Based Portal Segregation: Maintain discrete, secure login portals for Citizens, Department Officers, and Central Municipal Administrators.")
    add_bullet_point(doc, "Single-Origin Cloud Deployment: Deliver the entire platform as a unified web application requiring zero separate frontend or backend configurations.")
    
    add_subsection_heading(doc, "3.3 Expected Business Outcome")
    add_body_paragraph(doc, 
        "Successful deployment of SmartCivic delivers accelerated municipal response times, quantifiable department performance metrics, elimination of duplicate complaint logging, and higher civic satisfaction ratings across municipal wards.")
    
    add_subsection_heading(doc, "3.4 Project Deliverables")
    deliv_data = [
        ["01", "Citizen Web Portal", "Single-page web application featuring responsive dark-mode UI, registration, 3-step issue wizard, image dropzone, live tracking, and profile management."],
        ["02", "Department & Admin Command Hub", "Operational dispatch interface for department officers to inspect evidence, dispatch road/electrical crews, update statuses, and monitor zonal SLAs."],
        ["03", "Unified REST API & MongoDB Database", "High-performance Express backend delivering JWT authentication, triage engine, file streaming, and persistent document storage on MongoDB Atlas."]
    ]
    add_styled_table(doc, ["#", "Deliverable", "Description"], deliv_data, col_widths=[0.6, 2.2, 3.97])

    # =========================================================================
    # PAGE 5: 04. FUNCTIONAL REQUIREMENTS
    # =========================================================================
    start_new_page("FUNCTIONAL REQUIREMENTS")
    add_section_title(doc, "04", "Functional Requirements")
    
    add_subsection_heading(doc, "4.1 User Functionalities")
    add_bullet_point(doc, "User Registration & Authentication: Secure citizen onboarding with email validation, mobile number, residential address, and encrypted passwords.")
    add_bullet_point(doc, "User Profile Management: Real-time profile inspection, Aadhaar masking toggle, and persistent contact detail updates.")
    add_bullet_point(doc, "Citizen Dashboard: Overview of active grievances, SLA countdowns, priority indicators, and immediate action triggers.")
    add_bullet_point(doc, "3-Step Report Issue Wizard: Interactive form with category selection, priority tagging, GPS coordinate mapping, and evidence photo upload.")
    add_bullet_point(doc, "Isolated My Complaints: Secure personal grievance repository displaying strictly tickets submitted by the authenticated citizen.")
    add_bullet_point(doc, "Live Complaint Tracking: Public and authenticated search bar displaying interactive 5-stage progress timelines and inspector notes.")
    add_bullet_point(doc, "Real-Time Notification Feed: Automated notification center broadcasting ticket status changes, crew assignments, and official broadcasts.")
    
    add_subsection_heading(doc, "4.2 Admin Functionalities")
    add_bullet_point(doc, "Administrative & Officer Login: Discrete authentication gateways enforcing role verification (admin and dept_officer).")
    add_bullet_point(doc, "Zonal Command Dashboard: Real-time municipal metrics displaying total complaints, pending inspections, active crews, and 88% SLA rate.")
    add_bullet_point(doc, "Department Ticket Filtering: Instant filtering of tickets across municipal agencies (PWD, BESCOM, BBMP Sanitation, BWSSB, Traffic).")
    add_bullet_point(doc, "Status Transition & Crew Assignment: Capability to transition tickets across Submitted, Verified, Assigned, In Progress, and Resolved.")
    add_bullet_point(doc, "City-Wide Complaint Audit: Full administrative visibility into all citizen grievances and unassigned guest submissions across the city.")
    
    add_subsection_heading(doc, "4.3 Functional Requirement Details")
    fr_data = [
        ["FR-001", "Authentication & Role Guard", "JWT Bearer authentication strictly segregating Citizen, Department, and Admin access", "High"],
        ["FR-002", "Evidence Photo Upload Pipeline", "Multipart/form-data upload storing photos to disk/cloud with static URL rendering", "High"],
        ["FR-003", "Citizen Complaint Isolation", "Strict database query filtering ensuring citizen complaints are never exposed to peers", "High"],
        ["FR-004", "Live Milestone Action Timeline", "Timestamped 5-stage progress tracking with assigned engineer logs and remarks", "High"],
        ["FR-005", "Automated SLA Triage Engine", "Rule-based triage computing estimated resolution dates based on severity and department", "Medium"]
    ]
    add_styled_table(doc, ["ID", "Requirement", "Description", "Priority"], fr_data, col_widths=[1.0, 2.1, 2.87, 0.8])

    # =========================================================================
    # PAGE 6: 05. NON-FUNCTIONAL REQUIREMENTS (Part 1)
    # =========================================================================
    start_new_page("SYSTEM REQUIREMENTS")
    add_section_title(doc, "05", "Non-Functional Requirements")
    
    add_card_box(doc, "Performance", 
        "The system delivers sub-second API response times (<250ms for authenticated queries) and smooth 60fps client rendering. Client bundles are optimized using Tailwind utility compilation and React 18 virtual DOM diffing.")
    
    add_card_box(doc, "Security", 
        "End-to-end security is enforced via bcrypt password hashing (10 salt rounds), signed HMAC-SHA256 JWT tokens, HTTP 403 route protection on admin endpoints, and strict exclusion of credentials from version control.")
    
    add_card_box(doc, "Scalability", 
        "Stateless REST API architecture allows seamless horizontal scaling. Document persistence on MongoDB Atlas utilizes replica sets with automated sharding and indexing on ticketId, createdBy, and status.")
    
    add_card_box(doc, "Usability", 
        "Cyber-civic glassmorphic dark interface adhering to WCAG 2.1 contrast standards. Intuitive multi-step wizards, clear visual badges, and responsive layouts across mobile, tablet, and desktop viewports.")

    add_card_box(doc, "Availability", 
        "Engineered for 99.9% uptime leveraging cloud deployment on Render web services with automatic container health monitoring and MongoDB Atlas cloud multi-region failover.")

    add_card_box(doc, "Maintainability", 
        "Modular Model-View-Controller (MVC) directory structure, decoupled API route handlers, and a suite of 20 automated integration test cases ensuring regression-free maintenance.")

    add_subsection_heading(doc, "5.1 Hardware Requirements")
    add_bullet_point(doc, "Processor: Intel Core i3 / AMD Ryzen 3 or higher (2.0 GHz+ dual-core)")
    add_bullet_point(doc, "Memory (RAM): Minimum 4 GB RAM (8 GB recommended for concurrent builds)")
    add_bullet_point(doc, "Storage: Minimum 500 MB free disk space for runtime dependencies and assets")

    add_subsection_heading(doc, "5.2 Software Requirements")
    add_bullet_point(doc, "Operating System: Windows 10 / 11, macOS 12+, Ubuntu Linux 20.04 LTS+")
    add_bullet_point(doc, "Web Browser: Google Chrome 90+, Mozilla Firefox 88+, Microsoft Edge 90+, Safari 14+")

    # =========================================================================
    # PAGE 7: 05. NON-FUNCTIONAL REQUIREMENTS (Part 2 - Continued)
    # =========================================================================
    start_new_page("SYSTEM REQUIREMENTS")
    add_bullet_point(doc, "Development Environment: Visual Studio Code, Node.js v18.0.0+ (LTS), Python 3.10+ (Fallback engine)")
    add_bullet_point(doc, "Other Tools: MongoDB Atlas Cloud Cluster, Postman API Client, Git v2.40+, GitHub CLI (gh), Render Web Service Platform")
    
    add_card_box(doc, "System Compliance & Engineering Standards",
        "SmartCivic adheres to standard industry engineering guidelines including RESTful architectural constraints, standard HTTP status conventions (200, 201, 400, 401, 403, 404, 500), secure cookie/header transmission, and comprehensive Git commit conventions.")

    # =========================================================================
    # PAGE 8: 06. TECHNOLOGY STACK
    # =========================================================================
    start_new_page("TECHNOLOGY STACK")
    add_section_title(doc, "06", "Technology Stack")
    
    tech_data = [
        ["Frontend", "React 18, Tailwind CSS, Lucide Icons", "Single Page Application (SPA), responsive UI, glassmorphism design"],
        ["Backend", "Node.js, Express.js (Python HTTP fallback)", "RESTful API routing, middleware pipeline, static asset delivery"],
        ["Database", "MongoDB Atlas / Mongoose ODM", "Schema-based document persistence for users, complaints, notifications"],
        ["API", "RESTful JSON Architecture", "Standard HTTP protocol communication for client-server operations"],
        ["Authentication", "JSON Web Tokens (JWT), bcryptjs", "Cryptographic session tokens and salted password hashing"],
        ["Version Control", "Git / GitHub", "Source code versioning, branch management, collaborative tracking"],
        ["Deployment", "Render Cloud Web Service", "Production hosting with single-service unified static + API delivery"]
    ]
    add_styled_table(doc, ["Category", "Technology", "Purpose"], tech_data, col_widths=[1.5, 2.5, 2.77])
    
    add_subsection_heading(doc, "Development Tools")
    add_bullet_point(doc, "Visual Studio Code: Primary code editor with ESLint, Prettier, and Tailwind CSS IntelliSense.")
    add_bullet_point(doc, "Git & GitHub CLI: Version control management and automated repository deployment.")
    add_bullet_point(doc, "Postman: REST API development, endpoint validation, and automated test collection execution.")
    add_bullet_point(doc, "Python Test Runner: Automated regression test suites verifying 20 integration test cases (TC01-TC20).")

    # =========================================================================
    # PAGE 9: 07. SYSTEM ARCHITECTURE
    # =========================================================================
    start_new_page("SYSTEM ARCHITECTURE")
    add_section_title(doc, "07", "System Architecture")
    
    add_body_paragraph(doc, 
        "SmartCivic is architected as a modern 3-Tier Full-Stack Web Application delivering high throughput, decoupled component maintenance, and seamless single-origin cloud execution. The system unifies client presentation, server business logic, and cloud database persistence into an integrated lifecycle.")
    
    add_dashed_placeholder_box(doc, 
        "=========================================================================\n"
        "                  SMARTCIVIC SYSTEM ARCHITECTURE DIAGRAM\n"
        "=========================================================================\n\n"
        "  [ CLIENT PRESENTATION LAYER ]\n"
        "  React 18 SPA | Tailwind CSS | Glassmorphism | LocalStorage JWT Cache\n"
        "       │                                              ▲\n"
        "       │ HTTP Requests (JSON / FormData)              │ JSON Responses\n"
        "       ▼                                              │\n"
        "  [ APPLICATION & API LAYER (Express.js on Port 5000 / Render Cloud) ]\n"
        "  ├── Frontend Static Server (index.html, app.js, styles.css, /uploads)\n"
        "  ├── Authentication Middleware (JWT Verification & bcrypt hashing)\n"
        "  ├── Role Guard Interceptor (Citizen / Department / Municipal Admin)\n"
        "  ├── Complaint Management & AI SLA Triage Engine\n"
        "  └── Multer File Upload Stream Pipeline\n"
        "       │                                              ▲\n"
        "       │ Mongoose Queries                             │ Document Streams\n"
        "       ▼                                              │\n"
        "  [ CLOUD DATA PERSISTENCE LAYER (MongoDB Atlas M0 Cluster) ]\n"
        "  ├── Users Collection (Citizen, Dept Officer, Admin Documents)\n"
        "  ├── Complaints Collection (Grievances, Coordinates, SLA Timelines)\n"
        "  └── Notifications Collection (Real-Time Broadcast & Status Alerts)\n"
        "=========================================================================", height_inches=2.2)
    
    add_subsection_heading(doc, "7.1 Architecture Components")
    add_bullet_point(doc, "Frontend Layer: Single Page Application executing in the user's browser, managing state transitions, local token storage, dynamic view routing, and reactive UI updates.")
    add_bullet_point(doc, "Backend Layer: Express.js web server orchestrating RESTful endpoints, JWT authentication filters, Multer file upload pipelines, and automated SLA calculation rules.")
    add_bullet_point(doc, "Database Layer: MongoDB Atlas cloud database cluster maintaining structured collections for users, complaints, and notifications with indexed query performance.")
    add_bullet_point(doc, "API Layer: Standardized REST API exposing secure endpoints for authentication, complaint filing, tracking, and department status management.")
    add_bullet_point(doc, "External Services: Cloudinary / Render disk storage for complaint evidence photos and MongoDB Atlas multi-region cloud cluster.")
    
    add_subsection_heading(doc, "7.2 Data Flow")
    add_dashed_placeholder_box(doc, 
        "DATA FLOW: Citizen Client ➔ HTTP Request (Bearer JWT) ➔ Express Auth Middleware ➔ Route Controller ➔ Mongoose ODM ➔ MongoDB Atlas ➔ JSON Response ➔ React State Update", height_inches=0.8)

    # =========================================================================
    # PAGE 10: 08. SYSTEM WORKFLOW
    # =========================================================================
    start_new_page("SYSTEM WORKFLOW")
    add_section_title(doc, "08", "System Workflow")
    
    add_subsection_heading(doc, "8.1 Overall Application Workflow")
    add_dashed_placeholder_box(doc, 
        "[ Public Landing ] ➔ [ Login / Register Gateway ] ➔ [ JWT Role Verification ]\n"
        "      ├─ Citizen Role ➔ [ Citizen Dashboard: Report Wizard / My Complaints / Tracking ]\n"
        "      ├─ Dept Officer Role ➔ [ Department Portal: Filter Ward / Update Status / Dispatch Crew ]\n"
        "      └─ Admin Role ➔ [ Central Command: Zonal SLA Metrics / All City Complaints Audit ]", height_inches=1.0)
    
    add_subsection_heading(doc, "8.2 User Workflow")
    add_dashed_placeholder_box(doc, 
        "Citizen Registration/Login ➔ Open 3-Step Wizard ➔ Select Category & Coordinates ➔ Upload Photo Evidence ➔ Submit Grievance ➔ Automated SLA Triage ➔ Live Ticket Tracking ➔ Notification Receipt", height_inches=0.9)
    
    add_subsection_heading(doc, "8.3 Admin Workflow")
    add_dashed_placeholder_box(doc, 
        "Admin/Dept Login ➔ Review Filtered Department Wards ➔ Inspect Uploaded Photo & Coordinates ➔ Assign Field Officer ➔ Transition Status (Verified -> Assigned -> In Progress -> Resolved) ➔ Append Audit Log", height_inches=0.9)

    # =========================================================================
    # PAGE 11: 09. PROJECT MODULES
    # =========================================================================
    start_new_page("PROJECT MODULES")
    add_section_title(doc, "09", "Project Modules")
    
    add_card_box(doc, "9.1 Citizen Grievance & Reporting Module",
        "Purpose: Empowers authenticated residents to report neighborhood infrastructure defects with rich photographic evidence and precise geolocation coordinates.\n"
        "Features: 3-step reporting wizard, drag-and-drop image dropzone, category selector, GPS coordinate picker, and celebration confetti animation.\n"
        "Process: Validates form inputs, streams uploaded image to storage, generates unique ticket ID (e.g. SC-2026-0041), executes AI SLA triage, and records grievance.\n"
        "Input: Complaint title, civic category, priority rating, textual description, landmark/coordinates, and evidence photo file.\n"
        "Output: Persistent complaint document in MongoDB, unique ticket ID, isolated citizen list update, and real-time notification alert.")
    
    add_card_box(doc, "9.2 Department Dispatch & Status Management Module",
        "Purpose: Provides departmental field engineers (PWD, BESCOM, BBMP Sanitation, BWSSB) a dedicated portal to triage and resolve assigned civic complaints.\n"
        "Features: Pre-filtered departmental ticket streams, status transition controls, field officer assignment, and timestamped resolution remarks.\n"
        "Process: Authenticates officer JWT role, matches department jurisdiction, updates complaint lifecycle state, and logs audit timeline entries.\n"
        "Input: Selected ticket ID, target status (Verified, Assigned, In Progress, Resolved), assigned engineer name, and resolution remark.\n"
        "Output: Updated complaint record, appended milestone timeline log, and automated notification dispatched to the reporting citizen.")

    # =========================================================================
    # PAGE 12: 10. UI/UX DOCUMENTATION
    # =========================================================================
    start_new_page("UI / UX DOCUMENTATION")
    add_section_title(doc, "10", "UI/UX Documentation")
    
    add_subsection_heading(doc, "10.1 Design Overview")
    add_body_paragraph(doc, 
        "SmartCivic implements a modern Cyber-Civic design system characterized by deep slate backgrounds (#080C14, #0D121E), translucent glassmorphic cards with subtle cyan/purple/amber borders, glowing status indicators, and clean typography utilizing Outfit for headers and Plus Jakarta Sans for body elements. All components prioritize accessibility and responsive mobile performance.")
    
    add_subsection_heading(doc, "10.2 Navigation Structure")
    add_body_paragraph(doc, 
        "The user interface is driven by a single-page state router seamlessly rendering public views (Landing, Public Tracker, About, Contact), authentication gateways (Citizen Login, Admin Login, Department Login, Register), and authenticated role portals (Citizen Dashboard, Report Wizard, My Complaints, Complaint Detail, Track Complaints, Notifications, Profile, and Admin Dashboard).")
    
    add_subsection_heading(doc, "10.3 Page Documentation")
    add_body_paragraph(doc, "Page: Citizen Dashboard & Grievance Reporting Interface\nPurpose: Acts as the central interactive hub for registered citizens to file complaints, monitor ongoing ticket lifecycles, and review updates.")
    
    add_dashed_placeholder_box(doc, 
        "=========================================================================\n"
        "           INSERT PAGE SCREENSHOT: CITIZEN DASHBOARD & WIZARD\n"
        "=========================================================================\n"
        "  [ Active Complaints: 3 ]   [ Under Review: 1 ]   [ Resolved Issues: 2 ]\n"
        "  -----------------------------------------------------------------------\n"
        "  TICKET ID      CATEGORY         PRIORITY    STATUS        EST. RESOLVE\n"
        "  SC-2026-0041   Roads/Potholes   HIGH        In Progress   Jul 24, 2026\n"
        "  SC-2026-0038   Streetlights     MEDIUM      Resolved      Jul 19, 2026\n"
        "  SC-2026-0031   Waste/Sanitation HIGH        Verified      Jul 17, 2026\n"
        "=========================================================================", height_inches=1.8)
    
    add_body_paragraph(doc, "Figure 10.1 - SmartCivic Citizen Portal Interface", space_after=4)
    add_body_paragraph(doc, "Description: The interface displays real-time statistics (Active Complaints, Under Review, Resolved), recent ticket cards with priority badges, interactive upvote buttons, and quick navigation triggers to the 3-step complaint reporting wizard.")

    # =========================================================================
    # PAGE 13: 11. DATABASE DESIGN
    # =========================================================================
    start_new_page("DATABASE DESIGN")
    add_section_title(doc, "11. Database Design")
    
    add_subsection_heading(doc, "11.1 Database Overview")
    add_body_paragraph(doc, 
        "The database architecture is designed using MongoDB's document-oriented model, providing high throughput, flexible document structures, and relational linking via MongoDB ObjectIds. Three primary collections—users, complaints, and notifications—manage all civic data with indexing on email, ticketId, createdBy, and department.")
    
    add_subsection_heading(doc, "11.2 Database Tables / Collections")
    db_data = [
        ["users", "Stores citizen, department officer, and municipal admin profiles, roles, and hashed credentials", "_id (ObjectId)"],
        ["complaints", "Stores civic grievance records, image URLs, coordinates, department routing, and SLA timelines", "_id (ObjectId) / ticketId"],
        ["notifications", "Stores automated ticket status update alerts and municipal broadcast messages", "_id (ObjectId)"]
    ]
    add_styled_table(doc, ["Name", "Purpose", "Primary Key"], db_data, col_widths=[1.5, 3.77, 1.5])
    
    add_subsection_heading(doc, "11.3 ER Diagram")
    add_dashed_placeholder_box(doc, 
        "=========================================================================\n"
        "                     ENTITY RELATIONSHIP (ER) DIAGRAM\n"
        "=========================================================================\n"
        "  ┌───────────────────────┐          1:N         ┌───────────────────────┐\n"
        "  │         USER          │─────────────────────<│       COMPLAINT       │\n"
        "  ├───────────────────────┤ (createdBy)          ├───────────────────────┤\n"
        "  │ _id (ObjectId, PK)    │                      │ _id (ObjectId, PK)    │\n"
        "  │ name (String)         │                      │ ticketId (String, UQ) │\n"
        "  │ email (String, UQ)    │                      │ title (String)        │\n"
        "  │ password (Hash)       │                      │ category (String)     │\n"
        "  │ role (citizen/dept/adm│                      │ priority (String)     │\n"
        "  │ department (String)   │                      │ department (String)   │\n"
        "  │ mobile (String)       │                      │ status (String)       │\n"
        "  │ address (String)      │                      │ images (Array<String>)│\n"
        "  │ createdAt (Date)      │                      │ createdBy (FK -> User)│\n"
        "  └───────────────────────┘                      │ timeline (Array<Obj>) │\n"
        "              │                                  └───────────────────────┘\n"
        "              │ 1:N (user)                                   ▲\n"
        "              ▼                                              │ 1:1\n"
        "  ┌──────────────────────────────────────────────────────────┴───────────┐\n"
        "  │                             NOTIFICATION                             │\n"
        "  ├──────────────────────────────────────────────────────────────────────┤\n"
        "  │ _id (ObjectId, PK) | ticketId (String) | message (String) | user (FK)│\n"
        "  └──────────────────────────────────────────────────────────────────────┘\n"
        "=========================================================================", height_inches=2.2)

    # =========================================================================
    # PAGE 14: 12. API DOCUMENTATION
    # =========================================================================
    start_new_page("API DOCUMENTATION")
    add_section_title(doc, "12. API Documentation")
    
    add_body_paragraph(doc, "SmartCivic exposes a comprehensive RESTful JSON API secured with JWT Bearer tokens:")
    
    api_data = [
        ["/api/users/register", "POST", "Registers a new citizen with bcrypt password encryption"],
        ["/api/users/login", "POST", "Authenticates user credentials and returns signed JWT token"],
        ["/api/users/me", "GET", "Retrieves profile of currently authenticated user (Bearer token)"],
        ["/api/complaints", "GET", "Returns all complaints across the municipality (Admin / Zonal view)"],
        ["/api/complaints/user", "GET", "Returns complaints created strictly by authenticated citizen (Isolated)"],
        ["/api/complaints", "POST", "Creates new complaint with photo evidence and auto SLA triage"],
        ["/api/complaints/track/:id", "GET", "Public tracking endpoint returning live complaint timeline"],
        ["/api/admin/stats", "GET", "Returns zonal resolution statistics (Admin & Department only)"],
        ["/api/admin/complaints/:id/status", "PUT", "Updates complaint status, assigned officer, and timeline log"]
    ]
    add_styled_table(doc, ["Endpoint", "Method", "Purpose"], api_data, col_widths=[2.4, 0.9, 3.47])
    
    add_subsection_heading(doc, "API Request / Response")
    add_card_box(doc, "POST /api/complaints (Create Grievance with Evidence)",
        "Request Payload:\n"
        "{\n"
        '  "title": "Pothole on 5th Cross Main", "category": "Roads & Potholes", "priority": "High",\n'
        '  "description": "Deep asphalt depression causing risk", "location": "Indiranagar Zone C",\n'
        '  "images": ["/uploads/evidence-1788111446.jpg"]\n'
        "}\n\n"
        "Response (201 Created):\n"
        "{\n"
        '  "message": "Complaint registered successfully",\n'
        '  "complaint": { "ticketId": "SC-2026-0041", "status": "Submitted", "department": "PWD", "estResolution": "Jul 24, 2026" }\n'
        "}")

    # =========================================================================
    # PAGE 15: 13. IMPLEMENTATION DETAILS
    # =========================================================================
    start_new_page("IMPLEMENTATION")
    add_section_title(doc, "13. Implementation Details")
    
    add_subsection_heading(doc, "13.1 Project Structure")
    add_card_box(doc, "Directory Hierarchy",
        "smartcivic/\n"
        "├── index.html            # Main HTML5 SPA entry point\n"
        "├── app.js                # React 18 frontend engine & role router\n"
        "├── styles.css            # Cyber-civic dark theme & glassmorphic styles\n"
        "├── package.json          # Production dependencies & start script\n"
        "├── README.md             # Complete technical & setup documentation\n"
        "└── backend/\n"
        "    ├── server.js         # Unified Express server (static + REST API)\n"
        "    ├── config/db.js      # MongoDB Atlas connection manager\n"
        "    ├── models/           # Mongoose schemas (User, Complaint, Notification)\n"
        "    ├── routes/           # Express API route endpoints\n"
        "    ├── controllers/      # Business logic controllers\n"
        "    ├── middleware/       # JWT auth, role guards & upload pipeline\n"
        "    └── uploads/          # Stored evidence photos & attachments")
    
    add_subsection_heading(doc, "13.2 Frontend Implementation")
    add_body_paragraph(doc, 
        "The frontend is constructed with React 18 and Tailwind CSS, utilizing a centralized state engine in App() to manage view transitions, session persistence, active complaints, and notifications. Components employ conditional rendering to display citizen, department, and admin dashboards while maintaining reactive UI updates.")
    
    add_subsection_heading(doc, "13.3 Backend Implementation")
    add_body_paragraph(doc, 
        "The backend is powered by Node.js and Express.js, implementing clean MVC separation. The server serves both static frontend assets and RESTful API endpoints, utilizes Multer for multi-format image processing, and incorporates rule-based triage logic for automated SLA deadline computation.")
    
    add_subsection_heading(doc, "13.4 Database Integration")
    add_body_paragraph(doc, 
        "Database communication is managed via Mongoose ODM connecting to MongoDB Atlas. Schemas enforce strict type validation, default values, and relational references. An auto-seeding mechanism automatically populates demo users and sample complaints on fresh database initialization.")
    
    add_subsection_heading(doc, "13.5 Authentication")
    add_body_paragraph(doc, 
        "User authentication utilizes JSON Web Tokens (JWT) signed with HMAC-SHA256. Passwords are encrypted using bcrypt with a salt factor of 10. Express authMiddleware inspects Authorization: Bearer headers and enforces role guards for citizen, dept_officer, and admin privileges.")
    
    add_subsection_heading(doc, "13.6 Validation & Error Handling")
    add_body_paragraph(doc, 
        "Input payloads undergo server-side validation for required fields, email formats, and file MIME types. A centralized error handling middleware catches unhandled exceptions, returning standard HTTP error codes (400, 401, 403, 404, 500) and structured JSON error messages.")

    # =========================================================================
    # PAGE 16: 14. SECURITY IMPLEMENTATION
    # =========================================================================
    start_new_page("SECURITY")
    add_section_title(doc, "14. Security Implementation")
    
    add_bullet_point(doc, "Authentication: Signed JSON Web Tokens (JWT) with standard expiration and secure client-side storage.")
    add_bullet_point(doc, "Authorization: Multi-tier Role-Based Access Control segregating Citizen, Department, and Admin rights.")
    add_bullet_point(doc, "Password Protection: One-way bcrypt password hashing with cryptographic salt generation.")
    add_bullet_point(doc, "Input Validation: Strict server-side payload validation and file type checking (JPEG, PNG, WebP).")
    add_bullet_point(doc, "API Security: Protected REST endpoints verifying Bearer token authorization headers.")
    add_bullet_point(doc, "Role-Based Access Control: HTTP 403 Forbidden enforcement on administrative and dispatch APIs.")
    add_bullet_point(doc, "Secure Environment Variables: Complete isolation of secrets via .env and exclusion from Git.")
    add_bullet_point(doc, "Data Protection: Strict citizen complaint isolation preventing unauthorized access to peer complaints.")
    
    add_subsection_heading(doc, "Security Measures")
    add_body_paragraph(doc, 
        "1. Token-Based Security: All private API requests require a valid JWT passed in the Authorization header. Expired or tampered tokens are rejected with HTTP 401 Unauthorized.\n\n"
        "2. Role Guard Protection: Administrative APIs (/api/admin/stats, status updates) enforce explicit role checks. Citizen tokens attempting administrative access are blocked with HTTP 403 Forbidden.\n\n"
        "3. Prevention of Data Leakage (IDOR): Citizen complaint queries are scoped strictly to the authenticated user's MongoDB ObjectId (createdBy: user._id), preventing unauthorized visibility into other citizens' complaints.\n\n"
        "4. Secure Secret Management: Sensitive credentials including database connection strings and JWT secrets are managed exclusively through environment variables and excluded from Git via .gitignore.")

    # =========================================================================
    # PAGE 17: 15. TESTING & QUALITY ASSURANCE
    # =========================================================================
    start_new_page("TESTING & QA")
    add_section_title(doc, "15. Testing & Quality Assurance")
    
    add_subsection_heading(doc, "15.1 Testing Strategy")
    add_body_paragraph(doc, 
        "Quality assurance followed a rigorous test-driven approach including automated integration testing, unit API validation, security role enforcement checks, and end-to-end user flow verification using an automated Python test runner.")
    
    add_subsection_heading(doc, "15.2 Test Cases")
    tc_data = [
        ["TC-001", "Citizen Registration & JWT Token Issuance", "Returns HTTP 201 with signed token and user profile", "Token and profile generated accurately", "Pass"],
        ["TC-002", "Citizen Login with Bearer Token", "Authenticates credentials and restores user state", "Session state restored successfully", "Pass"],
        ["TC-003", "Role Guard Check on Admin Stats API", "Citizen token rejected with HTTP 403 Forbidden", "HTTP 403 Forbidden returned", "Pass"],
        ["TC-004", "Complaint Creation with Evidence Photo", "Stores image on disk and returns ticket ID", "Ticket created with /uploads/ image URI", "Pass"],
        ["TC-005", "Citizen Complaint Ownership Isolation", "Returns strictly complaints created by user", "Isolated user tickets verified", "Pass"],
        ["TC-006", "Department Officer Status Update & Timeline", "Updates status to In Progress and appends log", "Status and timeline updated", "Pass"]
    ]
    add_styled_table(doc, ["ID", "Test Scenario", "Expected Result", "Actual Result", "Status"], tc_data, col_widths=[0.8, 1.8, 1.8, 1.77, 0.6])
    
    add_subsection_heading(doc, "15.3 Bugs & Resolutions")
    bug_data = [
        ["Unassigned Complaint Leakage", "Unauthenticated submissions lacked explicit guest flag", "Set createdBy: null and isGuest: true with strict query filtering"],
        ["Babel Parsing Syntax Interruption", "Redundant closing JSX tags in login component", "Removed extra closing tags and validated with bracket parser"],
        ["Hardcoded Localhost API URL", "Client used static port 5000 in API requests", "Implemented dynamic relative /api detection for cloud hosting"]
    ]
    add_styled_table(doc, ["Issue", "Cause", "Resolution"], bug_data, col_widths=[1.8, 2.2, 2.77])

    # =========================================================================
    # PAGE 18: 16. PROJECT SCREENSHOTS
    # =========================================================================
    start_new_page("PROJECT OUTPUT")
    add_section_title(doc, "16. Project Screenshots")
    
    add_subsection_heading(doc, "16.1 Home / Landing Page")
    add_dashed_placeholder_box(doc, 
        "=========================================================================\n"
        "           INSERT SCREENSHOT: HOME / HERO LANDING PAGE\n"
        "  SmartCivic Civic Grievance & Resolution Platform | Live Zonal Activity\n"
        "  Active Reports: 1,420 | Resolved Issues: 1,180 | Live Resolution: 88%\n"
        "=========================================================================", height_inches=1.2)
    add_body_paragraph(doc, "Figure 16.1 - Home Page: Hero Landing with Live Activity Metrics", space_after=6)
    
    add_subsection_heading(doc, "16.2 Dashboard")
    add_dashed_placeholder_box(doc, 
        "=========================================================================\n"
        "           INSERT SCREENSHOT: CITIZEN DASHBOARD & TICKETS\n"
        "  Active Complaints | Priority Indicators | Live Action Timeline\n"
        "=========================================================================", height_inches=1.2)
    add_body_paragraph(doc, "Figure 16.2 - Citizen Dashboard: Grievances Table & Quick Actions", space_after=6)
    
    add_subsection_heading(doc, "16.3 Main Feature")
    add_dashed_placeholder_box(doc, 
        "=========================================================================\n"
        "           INSERT SCREENSHOT: 3-STEP REPORT ISSUE WIZARD\n"
        "  Step 1: Category & Priority | Step 2: Photo Dropzone & GPS | Step 3: Review\n"
        "=========================================================================", height_inches=1.2)
    add_body_paragraph(doc, "Figure 16.3 - Main Feature: 3-Step Report Issue Wizard with Evidence Upload", space_after=6)

    # =========================================================================
    # PAGE 19: 17. DEPLOYMENT DOCUMENTATION
    # =========================================================================
    start_new_page("DEPLOYMENT")
    add_section_title(doc, "17. Deployment Documentation")
    
    add_subsection_heading(doc, "17.1 Deployment Platform")
    add_body_paragraph(doc, 
        "Render Cloud Web Service (Node.js Environment) paired with MongoDB Atlas M0 Cloud Database Cluster.")
    
    add_subsection_heading(doc, "17.2 Deployment Process")
    add_body_paragraph(doc, 
        "1. Codebase Preparation: Unified Express backend to serve React static frontend files alongside REST API routes under a single origin.\n"
        "2. Version Control: Committed clean source code to GitHub repository (https://github.com/bharath20/SmartCivic.git) with sensitive files excluded.\n"
        "3. Render Service Setup: Created Web Service linked to repository with Build Command npm install and Start Command npm start.\n"
        "4. SSL & Domain Provisioning: Automatic SSL certificate generation delivering secure HTTPS single-URL access.")
    
    add_subsection_heading(doc, "17.3 Environment Configuration")
    add_body_paragraph(doc, 
        "The following environment variables are configured on Render:\n"
        "• NODE_ENV = production\n"
        "• MONGODB_URI = mongodb+srv://<username>:<password>@cluster0.xxxxx.mongodb.net/smartcivic?retryWrites=true&w=majority\n"
        "• JWT_SECRET = smartcivic_super_secret_jwt_key_2026\n"
        "• PORT = 5000 (Automatically assigned by Render runtime)")
    
    add_subsection_heading(doc, "17.4 Deployment Result")
    add_dashed_placeholder_box(doc, 
        "=========================================================================\n"
        "               INSERT DEPLOYMENT SCREENSHOT: RENDER DASHBOARD\n"
        "  Service: smartcivic-app | Status: Live | Runtime: Node.js 18 (Free Tier)\n"
        "  Public Deployed URL: https://smartcivic-app.onrender.com\n"
        "=========================================================================", height_inches=1.6)

    # =========================================================================
    # PAGE 20: 18. VERSION CONTROL & GITHUB
    # =========================================================================
    start_new_page("VERSION CONTROL")
    add_section_title(doc, "18. Version Control & GitHub")
    
    add_subsection_heading(doc, "18.1 Repository")
    add_card_box(doc, "GitHub Repository", "https://github.com/bharath20/SmartCivic")
    
    add_subsection_heading(doc, "18.2 Branch Structure")
    add_body_paragraph(doc, 
        "A clean trunk-based branching strategy was adopted. The main branch maintains the production-ready code with continuous integration and direct deployment to Render.")
    
    add_subsection_heading(doc, "18.3 Commit Strategy")
    add_body_paragraph(doc, 
        "Commits followed conventional semantic naming standards (feat:, fix:, chore:, test:, docs:) ensuring a transparent and auditable development history.")
    
    add_subsection_heading(doc, "18.4 Development Milestones")
    vc_data = [
        ["v1.0", "Core Frontend & Mock Engine", "Interactive React SPA with design slides, dark theme, and mock datasets"],
        ["v1.1", "Authentication & State Persistence", "Integrated JWT auth, profile editing, and MongoDB Mongoose schemas"],
        ["v1.2", "Evidence Upload & Citizen Isolation", "Multer multipart file streaming and strict complaint ownership separation"],
        ["v1.3", "3-Role Security & Unified Cloud Host", "Discrete Admin/Department login portals and Render single-server architecture"]
    ]
    add_styled_table(doc, ["Version", "Milestone", "Description"], vc_data, col_widths=[1.0, 2.4, 3.37])

    # =========================================================================
    # PAGE 21: 19. CHALLENGES & SOLUTIONS
    # =========================================================================
    start_new_page("PROJECT REVIEW")
    add_section_title(doc, "19. Challenges & Solutions")
    
    chal_data = [
        ["Multi-Role Portal Authentication", "Risk of unauthorized access to administrative controls", "Implemented discrete login views, backend JWT role verification, and client route guards"],
        ["Citizen Complaint Data Isolation", "Citizens initially saw all municipal complaints", "Restructured complaint schema with strict createdBy references and server-side filtering"],
        ["Evidence Image Persistence", "Cloud hosting instances have ephemeral file systems", "Implemented structured Multer upload pipelines with static serving and cloud-ready endpoints"],
        ["Single-Service Deployment Unification", "Separate frontend and backend servers caused CORS complexity", "Unified Express server to serve React SPA assets, SPA routing fallbacks, and REST APIs under one origin"]
    ]
    add_styled_table(doc, ["Challenge", "Impact", "Solution"], chal_data, col_widths=[1.8, 2.2, 2.77])

    # =========================================================================
    # PAGE 22: 20. PROJECT OUTCOME
    # =========================================================================
    start_new_page("PROJECT OUTCOME")
    add_section_title(doc, "20. Project Outcome")
    
    add_subsection_heading(doc, "20.1 Completed Features")
    add_bullet_point(doc, "3-Tier Role Portals (Citizen, Department Officer, Municipal Admin) with discrete authentication.")
    add_bullet_point(doc, "3-Step Interactive Issue Wizard with drag-and-drop evidence image upload and GPS mapping.")
    add_bullet_point(doc, "Isolated My Complaints dashboard displaying strictly authenticated user submissions.")
    add_bullet_point(doc, "Live Milestone Action Timeline with assigned officer updates and resolution remarks.")
    add_bullet_point(doc, "Department Dispatch Hub with jurisdiction filtering and crew assignment controls.")
    add_bullet_point(doc, "Central Municipal Command with real-time 88% SLA resolution metrics and audit trail.")
    
    add_subsection_heading(doc, "20.2 Final Project Status")
    add_status_badge(doc, "COMPLETED & FULLY DEPLOYED")
    
    add_subsection_heading(doc, "20.3 Key Achievements")
    add_bullet_point(doc, "20/20 automated integration test cases passing with zero failures.")
    add_bullet_point(doc, "Complete citizen complaint isolation with zero data leakage.")
    add_bullet_point(doc, "Sub-second API response latency (<250ms) across all endpoints.")
    add_bullet_point(doc, "Single-service public cloud deployment on Render connected to MongoDB Atlas.")
    
    add_subsection_heading(doc, "20.4 Expected Impact")
    add_body_paragraph(doc, 
        "SmartCivic establishes an accountable, digital-first municipal ecosystem that eliminates administrative opacity, accelerates civic repair turnaround times, and empowers citizens as active stakeholders in urban development.")

    # =========================================================================
    # PAGE 23: 21. FUTURE ENHANCEMENTS
    # =========================================================================
    start_new_page("PROJECT ROADMAP")
    add_section_title(doc, "21. Future Enhancements")
    
    add_body_paragraph(doc, "The following features and technological upgrades are planned for subsequent versions of SmartCivic:")
    
    add_bullet_point(doc, "AI Image Defect Verification: Integration of computer vision models to automatically detect pothole depth and garbage volume from uploaded photos prior to inspector dispatch.")
    add_bullet_point(doc, "WhatsApp & SMS Grievance Bot: Integration with messaging gateways allowing citizens to log complaints directly via messaging platforms without web app navigation.")
    add_bullet_point(doc, "Multilingual Regional Localization: Native language support for Kannada and Hindi to maximize accessibility across diverse demographic groups in Bengaluru.")
    add_bullet_point(doc, "IoT Telemetry Integration: Direct sensor feeds from smart streetlights and water flow meters for automated municipal ticket creation and dispatch.")

    # =========================================================================
    # PAGE 24: 22. CONCLUSION
    # =========================================================================
    start_new_page("CONCLUSION")
    add_section_title(doc, "22. Conclusion")
    
    add_body_paragraph(doc, 
        "SmartCivic successfully demonstrates a robust, scalable, and secure full-stack solution to modern urban grievance management. By integrating React 18, Express.js, and MongoDB Atlas into a unified single-application architecture, the platform bridges the communication gap between citizens and municipal authorities. With strict role-based access control, photographic evidence verification, citizen complaint isolation, and automated SLA triage, SmartCivic delivers a transparent, efficient, and community-centric civic technology ecosystem ready for real-world deployment.")
    
    add_subsection_heading(doc, "Key Takeaways")
    add_bullet_point(doc, "Technical Learning: Mastering full-stack MERN architecture, RESTful API design, JWT cryptographic security, and automated integration test automation.")
    add_bullet_point(doc, "Development Experience: Translating complex civic requirements and UI/UX design specifications into production-grade interactive React components.")
    add_bullet_point(doc, "Problem-Solving Experience: Successfully diagnosing and resolving real-world security vulnerabilities, complaint ownership isolation, and single-service deployment constraints.")
    add_bullet_point(doc, "Real-World Project Experience: Executing end-to-end cloud release management with Git, GitHub, Render, and MongoDB Atlas.")

    # =========================================================================
    # PAGE 25: 23. REFERENCES & APPENDIX
    # =========================================================================
    start_new_page("REFERENCES")
    add_section_title(doc, "23. References")
    
    references = [
        "1. React 18 Official Documentation & Hooks Reference — https://react.dev",
        "2. Express.js Application Framework & Middleware Pipeline Manual — https://expressjs.com",
        "3. MongoDB Atlas Database Administration & Mongoose ODM Guide — https://www.mongodb.com/docs",
        "4. JSON Web Token (JWT) Cryptographic Authentication Standard (RFC 7519) — https://jwt.io",
        "5. Tailwind CSS Utility-First Styling & Glassmorphism Guidelines — https://tailwindcss.com",
        "6. Code Morphicx 4th Batch Internship Project Guidelines & Academic Standards — Code Morphicx"
    ]
    for ref in references:
        add_body_paragraph(doc, ref, space_after=4)
        
    add_subsection_heading(doc, "Appendix")
    add_body_paragraph(doc, "Additional supporting technical materials including API endpoint specifications, Postman test collections, and zonal ward mapping:")
    add_dashed_placeholder_box(doc, 
        "=========================================================================\n"
        "               APPENDIX: SUPPORTING TECHNICAL MATERIAL\n"
        "=========================================================================\n"
        "  • Postman Automated Test Suite: SmartCivic_Postman_Collection.json (20 TCs)\n"
        "  • Postman Environment Secrets: SmartCivic_Postman_Environment.json\n"
        "  • Municipal Ward Jurisdiction: Bengaluru Municipal Zone C\n"
        "    (Covering Indiranagar, Ulsoor, MG Road, Cubbon Park, Domlur Wards)\n"
        "  • Verified Department Portals: PWD, BESCOM, BBMP Sanitation, BWSSB\n"
        "=========================================================================", height_inches=1.8)

    # =========================================================================
    # PAGE 26: BACK COVER PAGE
    # =========================================================================
    sec_back = doc.add_section(WD_SECTION.NEW_PAGE)
    sec_back.top_margin = Inches(1.5)
    sec_back.bottom_margin = Inches(1.0)
    sec_back.left_margin = Inches(0.75)
    sec_back.right_margin = Inches(0.75)
    sec_back.different_first_page_header_footer = True
    
    # Empty header/footer on back cover
    sec_back.header.is_linked_to_previous = False
    sec_back.header.paragraphs[0].text = ""
    sec_back.footer.is_linked_to_previous = False
    sec_back.footer.paragraphs[0].text = ""
    
    # Back Cover Logo
    if os.path.exists(logo_path):
        p_blogo = doc.add_paragraph()
        p_blogo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_blogo.paragraph_format.space_before = Pt(30)
        p_blogo.paragraph_format.space_after = Pt(16)
        p_blogo.add_run().add_picture(logo_path, width=Inches(1.2))
        
    p_bbrand = doc.add_paragraph()
    p_bbrand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bbrand.paragraph_format.space_before = Pt(0)
    p_bbrand.paragraph_format.space_after = Pt(6)
    r_bb = p_bbrand.add_run("CODE MORPHICX")
    r_bb.bold = True
    r_bb.font.size = Pt(22)
    r_bb.font.color.rgb = COLOR_NAVY_HEADER
    
    p_bsbar = doc.add_paragraph()
    p_bsbar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bsbar.paragraph_format.space_before = Pt(0)
    p_bsbar.paragraph_format.space_after = Pt(18)
    r_bsb = p_bsbar.add_run("━━━━━━")
    r_bsb.bold = True
    r_bsb.font.size = Pt(12)
    r_bsb.font.color.rgb = COLOR_PRIMARY_BLUE
    
    p_btitle = doc.add_paragraph()
    p_btitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_btitle.paragraph_format.space_before = Pt(0)
    p_btitle.paragraph_format.space_after = Pt(4)
    r_bt = p_btitle.add_run("Project Documentation")
    r_bt.bold = True
    r_bt.font.size = Pt(18)
    r_bt.font.color.rgb = COLOR_NAVY_HEADER
    
    p_bsub = doc.add_paragraph()
    p_bsub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bsub.paragraph_format.space_before = Pt(0)
    p_bsub.paragraph_format.space_after = Pt(35)
    r_bs = p_bsub.add_run("4th Batch Internship Program")
    r_bs.font.size = Pt(11)
    r_bs.font.color.rgb = COLOR_PRIMARY_BLUE
    
    p_bdesc = doc.add_paragraph()
    p_bdesc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bdesc.paragraph_format.space_before = Pt(0)
    p_bdesc.paragraph_format.space_after = Pt(40)
    r_bd = p_bdesc.add_run("This document represents the project assigned and developed as part of the Code Morphicx\nInternship Program.")
    r_bd.font.size = Pt(10)
    r_bd.font.color.rgb = COLOR_MUTED
    
    p_bdead = doc.add_paragraph()
    p_bdead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bdead.paragraph_format.space_before = Pt(0)
    p_bdead.paragraph_format.space_after = Pt(60)
    r_bdd = p_bdead.add_run("Documentation Submission Deadline\n28 August 2026")
    r_bdd.bold = True
    r_bdd.font.size = Pt(10)
    r_bdd.font.color.rgb = COLOR_TEXT_MAIN
    
    p_bcopy = doc.add_paragraph()
    p_bcopy.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bcopy.paragraph_format.space_before = Pt(40)
    p_bcopy.paragraph_format.space_after = Pt(0)
    r_bc = p_bcopy.add_run("© 2026 Code Morphicx. All Rights Reserved.")
    r_bc.font.size = Pt(9)
    r_bc.font.color.rgb = COLOR_MUTED

    # Output file
    output_filename = r"c:\Users\Bharath S\OneDrive\Desktop\smart civic\SmartCivic_Project_Documentation.docx"
    alt_filename = r"c:\Users\Bharath S\OneDrive\Desktop\smart civic\SmartCivic_Project_Documentation_Final.docx"
    
    saved_path = output_filename
    try:
        doc.save(output_filename)
    except PermissionError:
        doc.save(alt_filename)
        saved_path = alt_filename

    print(f"\n==================================================")
    print(f"[SUCCESS] 26-PAGE DOCUMENTATION GENERATED SUCCESSFULLY!")
    print(f"File Path: {saved_path}")
    print(f"Total Sections: {len(doc.sections)}")
    print(f"==================================================\n")

if __name__ == "__main__":
    build_smartcivic_documentation()
